"""
T.A.L.O.S. Document Processing Pipeline
Handles document ingestion, text extraction, and intelligent chunking

Supported formats:
- PDF (PyPDF2)
- DOCX (python-docx)
- TXT/Markdown
- Images with OCR (pytesseract)
"""

import os
import re
import io
import logging
import asyncio
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
import hashlib

import aiofiles
from PIL import Image

logger = logging.getLogger(__name__)


# =============================================================================
# ENUMS AND DATA CLASSES
# =============================================================================

class DocumentType(Enum):
    """Supported document types"""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MARKDOWN = "markdown"
    IMAGE = "image"
    UNKNOWN = "unknown"


class ChunkingStrategy(Enum):
    """Chunking strategies"""
    FIXED = "fixed"              # Fixed size chunks
    SENTENCE = "sentence"        # Sentence-based
    PARAGRAPH = "paragraph"      # Paragraph-based
    SEMANTIC = "semantic"        # Semantic boundaries
    HYBRID = "hybrid"            # Combined approach


@dataclass
class DocumentMetadata:
    """Metadata for a processed document"""
    filename: str
    document_type: DocumentType
    file_size: int
    page_count: int
    word_count: int
    char_count: int
    language: str = "en"
    title: Optional[str] = None
    author: Optional[str] = None
    created_date: Optional[str] = None
    checksum: str = ""
    extra: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "filename": self.filename,
            "document_type": self.document_type.value,
            "file_size": self.file_size,
            "page_count": self.page_count,
            "word_count": self.word_count,
            "char_count": self.char_count,
            "language": self.language,
            "title": self.title,
            "author": self.author,
            "created_date": self.created_date,
            "checksum": self.checksum,
            "extra": self.extra
        }


@dataclass
class TextChunk:
    """A chunk of extracted text"""
    chunk_id: int
    content: str
    start_char: int
    end_char: int
    page_number: Optional[int] = None
    section: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "chunk_id": self.chunk_id,
            "content": self.content,
            "start_char": self.start_char,
            "end_char": self.end_char,
            "page_number": self.page_number,
            "section": self.section,
            "metadata": self.metadata
        }


@dataclass
class ProcessedDocument:
    """Result of document processing"""
    document_id: int
    metadata: DocumentMetadata
    full_text: str
    chunks: List[TextChunk]
    success: bool
    error: Optional[str] = None
    processing_time_ms: float = 0.0

    def to_dict(self) -> Dict[str, Any]:
        return {
            "document_id": self.document_id,
            "metadata": self.metadata.to_dict(),
            "chunk_count": len(self.chunks),
            "success": self.success,
            "error": self.error,
            "processing_time_ms": self.processing_time_ms
        }


@dataclass
class ProcessorConfig:
    """Configuration for document processor"""
    # Chunking settings
    chunk_size: int = 500           # Target chunk size in characters
    chunk_overlap: int = 50         # Overlap between chunks
    min_chunk_size: int = 100       # Minimum chunk size
    max_chunk_size: int = 1000      # Maximum chunk size
    chunking_strategy: ChunkingStrategy = ChunkingStrategy.HYBRID

    # OCR settings
    ocr_enabled: bool = True
    ocr_language: str = "eng"
    ocr_dpi: int = 300

    # Processing settings
    extract_metadata: bool = True
    preserve_formatting: bool = False
    remove_headers_footers: bool = True

    def __post_init__(self):
        # Environment overrides
        self.chunk_size = int(os.getenv("CHUNK_SIZE", str(self.chunk_size)))
        self.chunk_overlap = int(os.getenv("CHUNK_OVERLAP", str(self.chunk_overlap)))
        self.ocr_enabled = os.getenv("OCR_ENABLED", "true").lower() == "true"
        self.ocr_language = os.getenv("OCR_LANGUAGE", self.ocr_language)


# =============================================================================
# TEXT EXTRACTORS
# =============================================================================

class PDFExtractor:
    """Extract text from PDF files"""

    @staticmethod
    async def extract(file_content: bytes, config: ProcessorConfig) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF.

        Args:
            file_content: Raw PDF bytes
            config: Processor configuration

        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(io.BytesIO(file_content))

            pages_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages_text.append(f"[Page {i + 1}]\n{text}")

            full_text = "\n\n".join(pages_text)

            # Extract metadata
            metadata = {}
            if reader.metadata:
                metadata = {
                    "title": reader.metadata.get("/Title", ""),
                    "author": reader.metadata.get("/Author", ""),
                    "subject": reader.metadata.get("/Subject", ""),
                    "creator": reader.metadata.get("/Creator", ""),
                }
            metadata["page_count"] = len(reader.pages)

            return full_text, metadata

        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise


class DOCXExtractor:
    """Extract text from DOCX files"""

    @staticmethod
    async def extract(file_content: bytes, config: ProcessorConfig) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from DOCX.

        Args:
            file_content: Raw DOCX bytes
            config: Processor configuration

        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_content))

            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # Check for headings
                    if para.style and para.style.name.startswith("Heading"):
                        paragraphs.append(f"\n## {para.text}\n")
                    else:
                        paragraphs.append(para.text)

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    table_text.append(row_text)
                if table_text:
                    paragraphs.append("\n" + "\n".join(table_text) + "\n")

            full_text = "\n\n".join(paragraphs)

            # Extract metadata
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }

            if doc.core_properties:
                metadata.update({
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                })

            return full_text, metadata

        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise


class ImageExtractor:
    """Extract text from images using OCR"""

    @staticmethod
    async def extract(file_content: bytes, config: ProcessorConfig) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from image using OCR.

        Args:
            file_content: Raw image bytes
            config: Processor configuration

        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        if not config.ocr_enabled:
            return "", {"ocr_skipped": True}

        try:
            import pytesseract

            # Open image
            image = Image.open(io.BytesIO(file_content))

            # Convert to RGB if necessary
            if image.mode != "RGB":
                image = image.convert("RGB")

            # Perform OCR
            text = pytesseract.image_to_string(
                image,
                lang=config.ocr_language,
                config=f"--dpi {config.ocr_dpi}"
            )

            metadata = {
                "image_size": image.size,
                "image_mode": image.mode,
                "ocr_language": config.ocr_language,
            }

            return text.strip(), metadata

        except Exception as e:
            logger.error(f"OCR extraction error: {e}")
            # Return empty on OCR failure (non-critical)
            return "", {"ocr_error": str(e)}


class TextExtractor:
    """Extract text from plain text files"""

    @staticmethod
    async def extract(file_content: bytes, config: ProcessorConfig) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from plain text file.

        Args:
            file_content: Raw text bytes
            config: Processor configuration

        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        try:
            # Try UTF-8 first, fallback to latin-1
            try:
                text = file_content.decode("utf-8")
            except UnicodeDecodeError:
                text = file_content.decode("latin-1")

            metadata = {
                "encoding": "utf-8" if "utf-8" in str(type(text)) else "latin-1",
                "line_count": text.count("\n") + 1,
            }

            return text, metadata

        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            raise


# =============================================================================
# CHUNKING STRATEGIES
# =============================================================================

class TextChunker:
    """Intelligent text chunking with multiple strategies"""

    def __init__(self, config: ProcessorConfig):
        self.config = config

        # Sentence-ending patterns
        self.sentence_pattern = re.compile(r'(?<=[.!?])\s+(?=[A-Z])')

        # Paragraph pattern
        self.paragraph_pattern = re.compile(r'\n\s*\n')

        # Section header pattern
        self.header_pattern = re.compile(r'^(?:#{1,6}\s+.+|[A-Z][^.!?]*:)\s*$', re.MULTILINE)

    def chunk(self, text: str, page_numbers: Optional[List[int]] = None) -> List[TextChunk]:
        """
        Chunk text using configured strategy.

        Args:
            text: Full text to chunk
            page_numbers: Optional page number mapping

        Returns:
            List of TextChunk objects
        """
        if self.config.chunking_strategy == ChunkingStrategy.FIXED:
            return self._fixed_chunk(text)
        elif self.config.chunking_strategy == ChunkingStrategy.SENTENCE:
            return self._sentence_chunk(text)
        elif self.config.chunking_strategy == ChunkingStrategy.PARAGRAPH:
            return self._paragraph_chunk(text)
        elif self.config.chunking_strategy == ChunkingStrategy.SEMANTIC:
            return self._semantic_chunk(text)
        else:  # HYBRID
            return self._hybrid_chunk(text)

    def _fixed_chunk(self, text: str) -> List[TextChunk]:
        """Fixed-size chunking with overlap"""
        chunks = []
        chunk_id = 0
        start = 0

        while start < len(text):
            end = start + self.config.chunk_size

            # Don't cut in the middle of a word
            if end < len(text):
                # Find nearest space
                space_pos = text.rfind(" ", start, end)
                if space_pos > start:
                    end = space_pos

            chunk_text = text[start:end].strip()

            if len(chunk_text) >= self.config.min_chunk_size:
                chunks.append(TextChunk(
                    chunk_id=chunk_id,
                    content=chunk_text,
                    start_char=start,
                    end_char=end
                ))
                chunk_id += 1

            # Move start with overlap
            start = end - self.config.chunk_overlap
            if start <= chunks[-1].start_char if chunks else 0:
                start = end  # Prevent infinite loop

        return chunks

    def _sentence_chunk(self, text: str) -> List[TextChunk]:
        """Sentence-based chunking"""
        sentences = self.sentence_pattern.split(text)
        chunks = []
        chunk_id = 0
        current_chunk = ""
        current_start = 0
        char_pos = 0

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            # Check if adding this sentence exceeds max size
            if len(current_chunk) + len(sentence) > self.config.max_chunk_size and current_chunk:
                chunks.append(TextChunk(
                    chunk_id=chunk_id,
                    content=current_chunk.strip(),
                    start_char=current_start,
                    end_char=char_pos
                ))
                chunk_id += 1
                current_chunk = sentence
                current_start = char_pos
            else:
                current_chunk += " " + sentence if current_chunk else sentence

            char_pos += len(sentence) + 1

        # Add remaining
        if current_chunk and len(current_chunk) >= self.config.min_chunk_size:
            chunks.append(TextChunk(
                chunk_id=chunk_id,
                content=current_chunk.strip(),
                start_char=current_start,
                end_char=char_pos
            ))

        return chunks

    def _paragraph_chunk(self, text: str) -> List[TextChunk]:
        """Paragraph-based chunking"""
        paragraphs = self.paragraph_pattern.split(text)
        chunks = []
        chunk_id = 0
        current_chunk = ""
        current_start = 0
        char_pos = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # If paragraph alone exceeds max, split it
            if len(para) > self.config.max_chunk_size:
                if current_chunk:
                    chunks.append(TextChunk(
                        chunk_id=chunk_id,
                        content=current_chunk.strip(),
                        start_char=current_start,
                        end_char=char_pos
                    ))
                    chunk_id += 1
                    current_chunk = ""

                # Split large paragraph using sentence chunking
                sub_chunks = self._sentence_chunk(para)
                for sub in sub_chunks:
                    sub.chunk_id = chunk_id
                    sub.start_char += char_pos
                    sub.end_char += char_pos
                    chunks.append(sub)
                    chunk_id += 1

                current_start = char_pos + len(para)

            # Check if adding this paragraph exceeds max size
            elif len(current_chunk) + len(para) > self.config.max_chunk_size and current_chunk:
                chunks.append(TextChunk(
                    chunk_id=chunk_id,
                    content=current_chunk.strip(),
                    start_char=current_start,
                    end_char=char_pos
                ))
                chunk_id += 1
                current_chunk = para
                current_start = char_pos
            else:
                current_chunk += "\n\n" + para if current_chunk else para

            char_pos += len(para) + 2

        # Add remaining
        if current_chunk and len(current_chunk) >= self.config.min_chunk_size:
            chunks.append(TextChunk(
                chunk_id=chunk_id,
                content=current_chunk.strip(),
                start_char=current_start,
                end_char=char_pos
            ))

        return chunks

    def _semantic_chunk(self, text: str) -> List[TextChunk]:
        """Semantic chunking based on section headers"""
        chunks = []
        chunk_id = 0

        # Find all section headers
        headers = list(self.header_pattern.finditer(text))

        if not headers:
            # Fall back to paragraph chunking
            return self._paragraph_chunk(text)

        # Process each section
        for i, match in enumerate(headers):
            section_start = match.start()
            section_end = headers[i + 1].start() if i + 1 < len(headers) else len(text)
            section_text = text[section_start:section_end].strip()
            section_header = match.group().strip()

            # If section is too large, sub-chunk it
            if len(section_text) > self.config.max_chunk_size:
                sub_chunks = self._paragraph_chunk(section_text)
                for sub in sub_chunks:
                    sub.chunk_id = chunk_id
                    sub.section = section_header
                    sub.start_char += section_start
                    sub.end_char += section_start
                    chunks.append(sub)
                    chunk_id += 1
            elif len(section_text) >= self.config.min_chunk_size:
                chunks.append(TextChunk(
                    chunk_id=chunk_id,
                    content=section_text,
                    start_char=section_start,
                    end_char=section_end,
                    section=section_header
                ))
                chunk_id += 1

        return chunks

    def _hybrid_chunk(self, text: str) -> List[TextChunk]:
        """
        Hybrid chunking strategy:
        1. Try semantic (section-based) first
        2. Fall back to paragraph for sections
        3. Use sentence splitting for overflow
        """
        # Check for clear section structure
        headers = list(self.header_pattern.finditer(text))

        if len(headers) >= 3:
            # Document has good structure, use semantic
            return self._semantic_chunk(text)

        # Check for paragraphs
        paragraphs = self.paragraph_pattern.split(text)
        if len(paragraphs) >= 3:
            return self._paragraph_chunk(text)

        # Fall back to sentence-based
        return self._sentence_chunk(text)


# =============================================================================
# DOCUMENT PROCESSOR
# =============================================================================

class DocumentProcessor:
    """
    Main document processing pipeline.

    Handles:
    - File type detection
    - Text extraction (PDF, DOCX, TXT, images)
    - Intelligent chunking
    - Metadata extraction
    """

    def __init__(self, config: Optional[ProcessorConfig] = None):
        self.config = config or ProcessorConfig()
        self.chunker = TextChunker(self.config)

        # File extension to type mapping
        self.extension_map = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".doc": DocumentType.DOCX,  # Will fail but try
            ".txt": DocumentType.TXT,
            ".md": DocumentType.MARKDOWN,
            ".markdown": DocumentType.MARKDOWN,
            ".png": DocumentType.IMAGE,
            ".jpg": DocumentType.IMAGE,
            ".jpeg": DocumentType.IMAGE,
            ".tiff": DocumentType.IMAGE,
            ".bmp": DocumentType.IMAGE,
            ".gif": DocumentType.IMAGE,
        }

        logger.info("Document Processor initialized")

    def detect_type(self, filename: str, content: Optional[bytes] = None) -> DocumentType:
        """
        Detect document type from filename or content.

        Args:
            filename: Original filename
            content: Optional file content for magic detection

        Returns:
            DocumentType enum
        """
        ext = Path(filename).suffix.lower()

        if ext in self.extension_map:
            return self.extension_map[ext]

        # Try magic detection if content provided
        if content:
            try:
                import magic
                mime = magic.from_buffer(content, mime=True)

                mime_map = {
                    "application/pdf": DocumentType.PDF,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocumentType.DOCX,
                    "text/plain": DocumentType.TXT,
                    "text/markdown": DocumentType.MARKDOWN,
                    "image/png": DocumentType.IMAGE,
                    "image/jpeg": DocumentType.IMAGE,
                }

                return mime_map.get(mime, DocumentType.UNKNOWN)
            except ImportError:
                pass

        return DocumentType.UNKNOWN

    async def process(
        self,
        document_id: int,
        filename: str,
        content: bytes
    ) -> ProcessedDocument:
        """
        Process a document through the full pipeline.

        Args:
            document_id: Unique document identifier
            filename: Original filename
            content: Raw file content

        Returns:
            ProcessedDocument with extracted text and chunks
        """
        import time
        start_time = time.time()

        try:
            # Detect document type
            doc_type = self.detect_type(filename, content)

            if doc_type == DocumentType.UNKNOWN:
                return ProcessedDocument(
                    document_id=document_id,
                    metadata=DocumentMetadata(
                        filename=filename,
                        document_type=doc_type,
                        file_size=len(content),
                        page_count=0,
                        word_count=0,
                        char_count=0
                    ),
                    full_text="",
                    chunks=[],
                    success=False,
                    error=f"Unsupported document type for file: {filename}",
                    processing_time_ms=(time.time() - start_time) * 1000
                )

            # Extract text based on type
            text, extra_metadata = await self._extract_text(content, doc_type)

            if not text.strip():
                return ProcessedDocument(
                    document_id=document_id,
                    metadata=DocumentMetadata(
                        filename=filename,
                        document_type=doc_type,
                        file_size=len(content),
                        page_count=extra_metadata.get("page_count", 0),
                        word_count=0,
                        char_count=0,
                        extra=extra_metadata
                    ),
                    full_text="",
                    chunks=[],
                    success=False,
                    error="No text could be extracted from document",
                    processing_time_ms=(time.time() - start_time) * 1000
                )

            # Clean text
            text = self._clean_text(text)

            # Chunk text
            chunks = self.chunker.chunk(text)

            # Build metadata
            checksum = hashlib.md5(content).hexdigest()
            metadata = DocumentMetadata(
                filename=filename,
                document_type=doc_type,
                file_size=len(content),
                page_count=extra_metadata.get("page_count", 1),
                word_count=len(text.split()),
                char_count=len(text),
                title=extra_metadata.get("title"),
                author=extra_metadata.get("author"),
                created_date=extra_metadata.get("created"),
                checksum=checksum,
                extra=extra_metadata
            )

            processing_time = (time.time() - start_time) * 1000

            logger.info(
                f"Processed {filename}: {len(chunks)} chunks, "
                f"{metadata.word_count} words, {processing_time:.1f}ms"
            )

            return ProcessedDocument(
                document_id=document_id,
                metadata=metadata,
                full_text=text,
                chunks=chunks,
                success=True,
                processing_time_ms=processing_time
            )

        except Exception as e:
            logger.error(f"Document processing error for {filename}: {e}")
            return ProcessedDocument(
                document_id=document_id,
                metadata=DocumentMetadata(
                    filename=filename,
                    document_type=DocumentType.UNKNOWN,
                    file_size=len(content),
                    page_count=0,
                    word_count=0,
                    char_count=0
                ),
                full_text="",
                chunks=[],
                success=False,
                error=str(e),
                processing_time_ms=(time.time() - start_time) * 1000
            )

    async def _extract_text(
        self,
        content: bytes,
        doc_type: DocumentType
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text based on document type"""

        if doc_type == DocumentType.PDF:
            return await PDFExtractor.extract(content, self.config)

        elif doc_type == DocumentType.DOCX:
            return await DOCXExtractor.extract(content, self.config)

        elif doc_type in (DocumentType.TXT, DocumentType.MARKDOWN):
            return await TextExtractor.extract(content, self.config)

        elif doc_type == DocumentType.IMAGE:
            return await ImageExtractor.extract(content, self.config)

        else:
            raise ValueError(f"Unsupported document type: {doc_type}")

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)

        # Remove common artifacts
        if self.config.remove_headers_footers:
            # Remove page numbers
            text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
            # Remove common header/footer patterns
            text = re.sub(r'\n.*?Page \d+ of \d+.*?\n', '\n', text)

        return text.strip()

    async def process_file(
        self,
        document_id: int,
        file_path: str
    ) -> ProcessedDocument:
        """
        Process a document from file path.

        Args:
            document_id: Unique document identifier
            file_path: Path to the file

        Returns:
            ProcessedDocument
        """
        async with aiofiles.open(file_path, "rb") as f:
            content = await f.read()

        filename = Path(file_path).name
        return await self.process(document_id, filename, content)


# =============================================================================
# FACTORY FUNCTION
# =============================================================================

_processor: Optional[DocumentProcessor] = None


def get_document_processor(config: Optional[ProcessorConfig] = None) -> DocumentProcessor:
    """Get or create document processor instance"""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor(config)
    return _processor


def reset_document_processor():
    """Reset document processor (for testing)"""
    global _processor
    _processor = None
